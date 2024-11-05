from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from faker import Faker
import os

fake = Faker("he_IL")


def add_hebrew_paragraph(document, text, is_heading=False, alignment=1):
    if is_heading:
        # Add a heading and align it to the right
        paragraph = document.add_heading(text, level=1)
    else:
        # Add a regular paragraph
        paragraph = document.add_paragraph(text)

    paragraph.alignment = alignment  # 2 represents alignment
    # Set the paragraph direction to RTL
    paragraph_element = paragraph._element
    paragraph_props = paragraph_element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    paragraph_props.append(bidi)
    return paragraph


def add_hebrew_paragraph_with_image(document, image_path):
    # Create a new paragraph and set the text
    paragraph = document.add_paragraph()
    run = paragraph.add_run("בברכה ")
    run = paragraph.add_run(fake.name())
    run = paragraph.add_run(f' מנכ"ל {fake.city()} ')

    # Set the paragraph direction to RTL (Right-to-Left)
    paragraph.alignment = 1  # 2 = Right alignment for Hebrew
    paragraph_element = paragraph._element
    paragraph_props = paragraph_element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    paragraph_props.append(bidi)

    # Add the image to the paragraph
    run.add_picture(image_path, width=Inches(2))  # Adjust the width as needed


def create_word_file(index):
    # Create a new Word document
    document = Document()

    # Add the date
    add_hebrew_paragraph(
        document, fake.date(pattern="%d/%m/%Y"), is_heading=False, alignment=2
    )

    # Add the title (heading) in RTL
    add_hebrew_paragraph(document, "אישור תושבות", is_heading=True, alignment=1)
    add_hebrew_paragraph(document, "", is_heading=False, alignment=0)
    add_hebrew_paragraph(document, "", is_heading=False, alignment=0)

    # Add the main text in RTL
    add_hebrew_paragraph(document, "לכל המעוניין,", is_heading=False, alignment=0)
    add_hebrew_paragraph(
        document,
        f" הריני לאשר כי {fake.name()} ת.ז. {fake.random_number(digits=9, fix_len=True)} ",
        is_heading=False,
        alignment=0,
    )
    add_hebrew_paragraph(
        document,
        f"מתגורר/ת בישוב {fake.city()} ברחוב: {fake.street_name()} {fake.random_int(min=1, max=100)} ",
        is_heading=False,
        alignment=0,
    )
    add_hebrew_paragraph(document, "", is_heading=False, alignment=0)
    add_hebrew_paragraph(
        document,
        f" מתאריך {fake.date(pattern='%d/%m/%Y')} ",
        is_heading=False,
        alignment=0,
    )
    add_hebrew_paragraph(document, "", is_heading=False, alignment=0)
    add_hebrew_paragraph(document, "", is_heading=False, alignment=0)
    add_hebrew_paragraph(document, "", is_heading=False, alignment=0)

    # Add an image
    add_hebrew_paragraph_with_image(document, image_path="./signature.PNG")

    # Save the document with a unique name
    document.save(f"word_files/אישור_תושבות_{index}.docx")


def create_multiple_word_files(num_files):
    # Create a directory to store the files if it doesn't exist
    if not os.path.exists("word_files"):
        os.makedirs("word_files")

    # Loop to create the specified number of Word files
    for i in range(1, num_files + 1):
        create_word_file(i)
        print(f"Created: אישור_תושבות_{i}.docx")


# Create 100 Word files
create_multiple_word_files(100)
